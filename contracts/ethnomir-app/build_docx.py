#!/usr/bin/env python3
"""Build a reference.docx styled in SF Pro Display (headings) / SF Pro Text (body),
   strict black-and-white (no colors, no horizontal rules), then convert v3.md to docx
   using pandoc, then post-process to:
   - force SF Pro fonts on every run (ascii, hAnsi, cs, eastAsia)
   - remove any bottom-border horizontal rule paragraphs
   - strip cell shading/fill and colored text
"""
import os, subprocess, sys
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

WORKDIR = "/home/claude/contract"
REFERENCE = os.path.join(WORKDIR, "reference.docx")
SRC_MD = os.path.join(WORKDIR, "Договор_ethnomir_app_v4.md")
OUT_DOCX = os.path.join(WORKDIR, "Договор_ethnomir_app_v4.docx")

HEADING_FONT = "SF Display"
BODY_FONT = "SF Text"

def set_font(run_element_rpr, font_name):
    """Set font on rPr element — on all four font attributes for full Unicode coverage."""
    # remove any existing rFonts
    for existing in run_element_rpr.findall(qn('w:rFonts')):
        run_element_rpr.remove(existing)
    r_fonts = OxmlElement('w:rFonts')
    r_fonts.set(qn('w:ascii'), font_name)
    r_fonts.set(qn('w:hAnsi'), font_name)
    r_fonts.set(qn('w:cs'), font_name)
    r_fonts.set(qn('w:eastAsia'), font_name)
    run_element_rpr.insert(0, r_fonts)

def set_style_font(style, font_name, size_pt=None, bold=None):
    font = style.font
    font.name = font_name
    if size_pt is not None:
        font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    # Colour: strictly black
    font.color.rgb = RGBColor(0, 0, 0)
    # Drill into XML to set rFonts on all scripts
    rpr = style.element.get_or_add_rPr()
    set_font(rpr, font_name)
    # also clear color xml
    for c in rpr.findall(qn('w:color')):
        rpr.remove(c)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '000000')
    rpr.append(color_el)

def build_reference_docx():
    doc = Document()

    # Default (Normal) — SF Pro Text 11pt
    normal = doc.styles['Normal']
    set_style_font(normal, BODY_FONT, size_pt=11, bold=False)

    # Heading 1 — SF Pro Display 22pt Bold
    set_style_font(doc.styles['Heading 1'], HEADING_FONT, size_pt=22, bold=True)
    # Heading 2 — SF Pro Display 16pt Bold
    set_style_font(doc.styles['Heading 2'], HEADING_FONT, size_pt=16, bold=True)
    # Heading 3 — SF Pro Display 13pt Bold
    set_style_font(doc.styles['Heading 3'], HEADING_FONT, size_pt=13, bold=True)
    # Heading 4 — SF Pro Display 12pt Bold
    if 'Heading 4' in [s.name for s in doc.styles]:
        set_style_font(doc.styles['Heading 4'], HEADING_FONT, size_pt=12, bold=True)

    # Title
    try:
        set_style_font(doc.styles['Title'], HEADING_FONT, size_pt=26, bold=True)
    except KeyError:
        pass

    # Table grid / Normal Table — neutral, black text
    for style_name in ['Table Grid', 'Normal Table']:
        try:
            st = doc.styles[style_name]
            set_style_font(st, BODY_FONT, size_pt=10, bold=False)
        except KeyError:
            pass

    # Page setup: US Letter analog, 1 inch margins (but A4 is fine for ru doc — keep A4)
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(1.5)

    doc.save(REFERENCE)
    print(f"[ok] reference.docx built at {REFERENCE}")

def run_pandoc():
    cmd = [
        "pandoc", SRC_MD,
        "--from", "markdown+raw_attribute+pipe_tables+backtick_code_blocks",
        "--to", "docx",
        "--reference-doc", REFERENCE,
        "-o", OUT_DOCX,
    ]
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.returncode != 0:
        print("pandoc stderr:", result.stderr)
        sys.exit(1)
    print("[ok] pandoc conversion done")

def post_process():
    """Force SF Pro fonts on every run, strip horizontal rule paragraphs,
       strip color overrides and cell shading, and fix table layout."""
    from docx import Document as D
    doc = D(OUT_DOCX)

    # helpers
    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def force_font_on_run(run):
        rpr = run._element.get_or_add_rPr()
        # decide which font: if run is inside a heading paragraph — SF Pro Display; else SF Pro Text
        # simpler: if any bold/size suggests heading, use Display. We'll set on parent paragraph level later.
        set_font(rpr, BODY_FONT)
        # strip color override (force black)
        for c in rpr.findall(qn('w:color')):
            rpr.remove(c)
        color_el = OxmlElement('w:color')
        color_el.set(qn('w:val'), '000000')
        rpr.append(color_el)
        # strip highlight/shading on run
        for sh in rpr.findall(qn('w:shd')):
            rpr.remove(sh)
        for hi in rpr.findall(qn('w:highlight')):
            rpr.remove(hi)

    def force_font_on_heading_run(run):
        rpr = run._element.get_or_add_rPr()
        set_font(rpr, HEADING_FONT)
        for c in rpr.findall(qn('w:color')):
            rpr.remove(c)
        color_el = OxmlElement('w:color')
        color_el.set(qn('w:val'), '000000')
        rpr.append(color_el)
        for sh in rpr.findall(qn('w:shd')):
            rpr.remove(sh)
        for hi in rpr.findall(qn('w:highlight')):
            rpr.remove(hi)

    # Paragraphs
    hr_paragraphs_to_remove = []
    for p in doc.paragraphs:
        style_name = (p.style.name if p.style else "") or ""
        is_heading = style_name.startswith("Heading") or style_name == "Title"

        # Remove paragraph-level bottom-border horizontal rules (pandoc --- produces these)
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                # if paragraph is empty or only whitespace and has a bottom border, mark for removal
                text = (p.text or "").strip()
                if not text:
                    hr_paragraphs_to_remove.append(p._element)
                    continue
                # remove the border itself but keep paragraph
                pPr.remove(pBdr)

        # Apply fonts on runs
        for run in p.runs:
            if is_heading:
                force_font_on_heading_run(run)
            else:
                force_font_on_run(run)

    for el in hr_paragraphs_to_remove:
        el.getparent().remove(el)

    # Tables — strip shading, set font on cells, fix widths and layout
    # Page content width for A4 with our margins: 21cm - 2.5 - 1.5 = 17cm = ~9640 DXA
    # We'll use 9200 DXA to be safe.
    TOTAL_DXA = 9200

    for table in doc.tables:
        tbl_el = table._element

        # remove table-level shading
        for tblPr in tbl_el.findall(qn('w:tblPr')):
            for sh in tblPr.findall(qn('w:shd')):
                tblPr.remove(sh)
            # remove tblStyle (points to undefined "Table" style → confuses renderer)
            for ts in tblPr.findall(qn('w:tblStyle')):
                tblPr.remove(ts)

        # 1. Compute column widths.
        # Strategy: look at max content length per column, but use log-scale
        # (short strings shouldn't get crushed; long strings shouldn't eat everything).
        # Then enforce a hard minimum width so words fit without letter-wrapping.
        import math
        rows = table.rows
        if not rows:
            continue
        ncols = len(rows[0].cells)

        # max length per column (capped at 120)
        col_max_len = [1] * ncols
        # also track longest single word per column (for hard minimum)
        col_longest_word = [1] * ncols
        for row in rows:
            for i, cell in enumerate(row.cells):
                if i < ncols:
                    txt = (cell.text or "").strip()
                    col_max_len[i] = max(col_max_len[i], min(len(txt), 120))
                    for w in txt.split():
                        col_longest_word[i] = max(col_longest_word[i], len(w))

        # log-scaled weights
        weights = [math.log(L + 2) for L in col_max_len]
        total_w = sum(weights)
        # proportional widths
        col_widths = [int(round(TOTAL_DXA * w / total_w)) for w in weights]

        # Enforce hard minimum ~ longest_word_chars * 120 dxa (each char ~ 6pt ≈ 120 dxa)
        # but not less than 900 dxa (~1.5cm)
        MIN_BASE_DXA = 900
        CHAR_DXA = 110
        min_widths = [max(MIN_BASE_DXA, col_longest_word[i] * CHAR_DXA) for i in range(ncols)]

        # Enforce minimums by raising narrow columns and shrinking wide ones proportionally.
        for _ in range(5):
            deficits = [max(0, min_widths[i] - col_widths[i]) for i in range(ncols)]
            total_deficit = sum(deficits)
            if total_deficit == 0:
                break
            # find columns with surplus (col_widths > min_widths by > 100)
            surpluses = [max(0, col_widths[i] - min_widths[i] - 100) for i in range(ncols)]
            total_surplus = sum(surpluses)
            if total_surplus == 0:
                break
            # proportionally take from surplus, give to deficit
            take = min(total_deficit, total_surplus)
            for i in range(ncols):
                if deficits[i] > 0:
                    col_widths[i] += int(take * deficits[i] / total_deficit)
                if surpluses[i] > 0:
                    col_widths[i] -= int(take * surpluses[i] / total_surplus)

        # Normalize to exactly TOTAL_DXA
        drift = TOTAL_DXA - sum(col_widths)
        col_widths[-1] += drift

        # 2. Apply to tblGrid
        grid = tbl_el.find(qn('w:tblGrid'))
        if grid is not None:
            for i, gridCol in enumerate(grid.findall(qn('w:gridCol'))):
                if i < ncols:
                    gridCol.set(qn('w:w'), str(col_widths[i]))

        # 3. Apply to tblPr: tblW type=dxa, tblLayout type=fixed
        tblPr = tbl_el.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl_el.insert(0, tblPr)
        # tblW
        for old in tblPr.findall(qn('w:tblW')):
            tblPr.remove(old)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:type'), 'dxa')
        tblW.set(qn('w:w'), str(TOTAL_DXA))
        tblPr.append(tblW)
        # tblLayout = fixed (prevents auto-resize chaos)
        for old in tblPr.findall(qn('w:tblLayout')):
            tblPr.remove(old)
        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

        # tblBorders — light grid so the structure reads as a table
        for old in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(old)
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{border_name}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')        # 0.5pt
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), '888888')  # light gray
            tblBorders.append(b)
        tblPr.append(tblBorders)

        # 4. Apply widths and fonts on each cell
        for row in rows:
            for i, cell in enumerate(row.cells):
                if i >= ncols:
                    continue
                tc = cell._element
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                # remove shading
                for sh in tcPr.findall(qn('w:shd')):
                    tcPr.remove(sh)
                # set tcW
                for old in tcPr.findall(qn('w:tcW')):
                    tcPr.remove(old)
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:type'), 'dxa')
                tcW.set(qn('w:w'), str(col_widths[i]))
                tcPr.append(tcW)
                # fix fonts inside cell
                for p in cell.paragraphs:
                    # strip pStyle references to undefined "Compact" style
                    pPr = p._element.find(qn('w:pPr'))
                    if pPr is not None:
                        for pStyle in pPr.findall(qn('w:pStyle')):
                            val = pStyle.get(qn('w:val'))
                            if val in ('Compact', 'Table', 'TableParagraph'):
                                pPr.remove(pStyle)
                    for run in p.runs:
                        force_font_on_run(run)

    doc.save(OUT_DOCX)
    print(f"[ok] post-processing complete — {OUT_DOCX}")
    print(f"     removed {len(hr_paragraphs_to_remove)} horizontal-rule paragraphs")

if __name__ == "__main__":
    build_reference_docx()
    run_pandoc()
    post_process()
    print("\nDONE.")
